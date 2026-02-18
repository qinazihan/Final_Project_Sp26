#!/usr/bin/env python3
"""
Title: MBA Strategy Agent v4 — Per-Topic Debate Loop + 3 Human Gates

A multi-agent business strategy assistant using LangGraph with ChatOpenAI
routed through OpenRouter. v4 introduces per-topic debate (Research & Propose
⇄ Critic) with human approval per topic, plus a final Human Gate 3.

All config and scenario inputs are loaded from config.yaml.

Pipeline:
  START → intake → clarify → [Gate 1] → plan_topics →
  ┌→ research_and_propose ⇄ topic_critic → [Gate 2] → (next topic)─┐
  └──────────────────────────────────────────────────────────────────┘
  → synthesizer → action_plan_90d → [Gate 3] → END

Input:  input_query from config.yaml or HumanMessage
Output: MBA-course/results/ (JSON + markdown report)
"""

# %% [markdown]
# # MBA Strategy Agent v4 — Per-Topic Debate + 3 Human Gates
#
# **Stack:** LangGraph + ChatOpenAI + OpenRouter
#
# **What's new in v4:**
# - Per-topic debate loop: Research & Propose ⇄ Critic converges before human review
# - Human Gate 2 fires per-topic (inside the loop)
# - New Human Gate 3 for final approval of the synthesized report
# - All config + scenario moved to config.yaml
#
# **Pipeline (11 nodes, 3 interrupt gates, 5 routing functions):**
# ```
# START → intake → clarify → Gate 1 → plan_topics →
# [research_and_propose ⇄ topic_critic → Gate 2] per topic →
# synthesizer → action_plan_90d → Gate 3 → END
# ```

# %% Setup and Imports
import os
import json
import operator
import textwrap
from pathlib import Path
from datetime import datetime
from typing import Annotated, Literal

import warnings
warnings.filterwarnings("ignore", message="Pydantic serializer warnings")

import yaml
from pydantic import BaseModel, Field
from typing_extensions import TypedDict

from langchain_openai import ChatOpenAI
from langchain_core.messages import AIMessage, HumanMessage, SystemMessage
from langgraph.graph import StateGraph, START, END, MessagesState
from langgraph.types import interrupt

from dotenv import load_dotenv
load_dotenv(os.path.join(os.path.dirname(__file__), ".env"))


# %% Load Configuration from YAML
_config_path = Path(os.getenv("MBA_CONFIG_PATH", Path(__file__).resolve().parent / "config.yaml"))
with open(_config_path) as _f:
    CFG = yaml.safe_load(_f)

# API
OPENROUTER_BASE_URL = CFG["openrouter_base_url"]
OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY", "")

# Agent configs
AGENTS = CFG["agents"]

# Agent behavior
AUTO_APPROVE = os.getenv("MBA_AUTO_APPROVE", str(CFG["auto_approve"])).lower() == "true"
MAX_CLARIFY_ROUNDS = CFG["max_clarify_rounds"]
MAX_RESEARCH_TOPICS = CFG["max_research_topics"]
MAX_WEB_SEARCH_CT = CFG["max_web_search_ct"]
MAX_DEBATE_ROUNDS = CFG["max_debate_rounds"]
MAX_HUMAN_REVISION_ON_PROPOSAL = CFG["max_human_revision_on_proposal"]
MAX_HUMAN_REVISION_ON_PLAN = CFG["max_human_revision_on_plan"]

# Context budget — sized for the smallest model (GPT-4o: 128K tokens ≈ 400K chars usable)
MAX_CONTEXT_CHARS = 300_000

# Default topics & scenario
INPUT_QUERY = CFG["input_query"]

print("=" * 80)
print("SETUP COMPLETE — MBA Agent v4")
print("=" * 80)

# Tavily web search (required)
from tavily import TavilyClient
_tavily_key = os.getenv("TAVILY_API_KEY", "")
if not _tavily_key:
    raise RuntimeError("TAVILY_API_KEY is required. Set it in .env or as an environment variable.")
_tavily = TavilyClient(api_key=_tavily_key)

print(f"  Web search: Tavily")

# %% Configuration Summary
print("=" * 80)
print("CONFIGURATION")
print("=" * 80)
print(f"  Auto-approve:      {AUTO_APPROVE}")
print(f"  Max debate rounds: {MAX_DEBATE_ROUNDS}")
print(f"  Max topics:        {MAX_RESEARCH_TOPICS}")
print(f"  Query:             {INPUT_QUERY}")

# %% Initialize LLMs — one per agent role
# Models that don't support max_tokens (use max_completion_tokens or omit entirely)
_NO_MAX_TOKENS_MODELS = {
    "openai/o1", "openai/o3", "openai/o3-mini", "openai/o4-mini",
    "openai/gpt-5.1", "openai/gpt-5.2", "openai/gpt-5.2-pro",
    "deepseek/deepseek-r1",
}

def _make_llm(agent_name: str, max_tokens: int = 4096) -> ChatOpenAI:
    cfg = AGENTS[agent_name]
    model_id = cfg["model"]
    kwargs = dict(
        model=model_id,
        base_url=OPENROUTER_BASE_URL,
        api_key=OPENROUTER_API_KEY,
        max_retries=3,
    )
    # Reasoning / frontier models don't support max_tokens
    if model_id not in _NO_MAX_TOKENS_MODELS:
        kwargs["max_tokens"] = max_tokens
    if "temperature" in cfg:
        kwargs["temperature"] = cfg["temperature"]
    return ChatOpenAI(**kwargs)

llm_intake = _make_llm("intake")        # Intake, Clarify, Plan Topics
llm_researcher = _make_llm("researcher") # Research & Propose
llm_critic = _make_llm("critic")         # Critic
llm_synthesizer = _make_llm("synthesizer", max_tokens=8192)  # Synthesizer, Action Plan

print("=" * 80)
print("LLMs INITIALIZED")
print("=" * 80)
for name, cfg in AGENTS.items():
    temp_str = f" (temp={cfg['temperature']})" if "temperature" in cfg else ""
    print(f"  {name:12s} → {cfg['model']}{temp_str}")

# %% Pydantic Output Models
class IntakeOutput(BaseModel):
    """Output of the Intake node — parses chat message into structured fields."""
    user_query: str = Field(description="The core business question")
    country_or_market: str = Field(default="Not specified", description="Country or market mentioned")
    product_idea: str = Field(default="Not specified", description="Product or service idea")
    target_customer: str = Field(default="Not specified", description="Target customer segment")
    budget_range: str = Field(default="Not specified", description="Budget mentioned")
    time_horizon: str = Field(default="Not specified", description="Time horizon mentioned")
    risk_tolerance: str = Field(default="Not specified", description="Risk tolerance mentioned")
    constraints: str = Field(default="None", description="Any constraints mentioned")

class ClarifyOutput(BaseModel):
    """Output of the Clarify Problem node."""
    problem_framing: str = Field(description="Clear 2-3 sentence problem statement")
    constraints_noted: str = Field(description="Key constraints identified")
    questions: list[str] = Field(default_factory=list, description="Questions for the user about unclear or missing aspects. Empty list if everything is clear.")

class TopicsOutput(BaseModel):
    """Output of the Plan Research Topics node."""
    topics: list[str] = Field(description="List of research topics to investigate")

class FindingSource(BaseModel):
    title: str = ""
    url: str = ""

class Finding(BaseModel):
    claim: str = Field(description="Specific research finding")
    confidence: str = Field(description="high, medium, or low")
    sources: list[FindingSource] = Field(default_factory=list)

class TopicProposalOutput(BaseModel):
    """Output of the Research & Propose node — per-topic."""
    topic: str = Field(description="The topic researched")
    findings: list[Finding] = Field(description="List of findings with confidence")
    summary: str = Field(description="2-3 sentence summary")
    proposal: str = Field(description="Strategic proposal for this topic")
    key_recommendation: str = Field(description="One-line recommendation")
    gap_responses: list[str] = Field(default_factory=list, description="On revision rounds: one response per critic gap, explaining how it was addressed or why it cannot be. Empty list on first round.")

class CriticOutput(BaseModel):
    """Output of the per-topic Critic node."""
    assessment: str = Field(description="Assessment of the proposal quality")
    gaps: list[str] = Field(description="Identified gaps or weaknesses that can be fixed. Empty list if converged.")
    converged: bool = Field(description="True if proposal is good enough for human review")
    revision_guidance: str = Field(description="If not converged, what to fix next round. Empty string if converged.")
    limitations: list[str] = Field(default_factory=list, description="Remaining limitations that cannot be resolved via web research. Only populated when converged.")

print("Pydantic models defined (8 output schemas)")

# %% State Schema
class MBAAgentState(MessagesState):
    # Inherits: messages: Annotated[list[AnyMessage], add_messages]
    # User inputs
    user_query: str
    country_or_market: str
    product_idea: str
    target_customer: str
    budget_range: str
    time_horizon: str
    risk_tolerance: str
    constraints: str
    # Intake agent conversation memory (cleared after plan_research_topics)
    intake_chat_history: list[dict]
    # Clarification
    problem_framing: str
    constraints_noted: str
    questions: list[str]
    # Research topics
    research_topics: list[str]
    # Handoff to researcher (final distilled output from intake phase)
    research_brief: str
    current_topic_idx: int
    # Per-topic debate state (overwritten each round)
    current_debate_round: int
    current_topic_proposal: str
    current_topic_critique: str
    current_topic_limitations: list[str]
    debate_converged: bool
    # Debate conversation history (researcher ⇄ critic + human feedback, cleared per topic)
    debate_history: list[dict]
    # Action plan conversation history (action_plan ⇄ human feedback at Gate 3)
    action_plan_history: list[dict]
    # Approved topics (append-only accumulator)
    approved_topics: Annotated[list[dict], operator.add]
    # Human feedback
    human_feedback_1: str
    human_feedback_2: str
    human_feedback_3: str
    # Control
    clarify_round: int
    proposal_revision_round: int
    plan_revision_round: int
    status: str
    # Final outputs
    recommendation: str
    action_plan: str
    final_output: str

print("State schema defined")

# %% Helpers
def _msg_text(msg) -> str:
    """Extract plain text from a message's content."""
    c = msg.content
    if isinstance(c, str):
        return c.strip()
    if isinstance(c, list):
        parts = []
        for block in c:
            if isinstance(block, str):
                parts.append(block)
            elif isinstance(block, dict) and "text" in block:
                parts.append(block["text"])
        return " ".join(parts).strip()
    return str(c).strip()

def _truncate(text: str, max_chars: int = MAX_CONTEXT_CHARS, label: str = "text") -> str:
    """Truncate text and print a warning if it was shortened."""
    if len(text) > max_chars:
        print(f"  [WARNING] {label} truncated: {len(text)} → {max_chars} chars")
        return text[:max_chars]
    return text

def web_search(query: str, max_results: int = 3) -> list[dict]:
    """Search the web via Tavily. Raises on failure."""
    results = _tavily.search(query=query, max_results=max_results)
    return [
        {"title": r.get("title", ""), "url": r.get("url", ""), "content": r.get("content", "")}
        for r in results.get("results", [])
    ]

def _build_debate_context(debate_history: list[dict], max_chars: int = MAX_CONTEXT_CHARS) -> str:
    """Build debate history context string, truncating older entries if too long."""
    if not debate_history:
        return ""
    entries = []
    char_count = 0
    for entry in reversed(debate_history):
        role = entry.get("role", "unknown")
        content = entry.get("content", "")
        content = _truncate(content, label=f"debate entry ({role})")
        entry_text = f"\n\n[{role.upper()}]:\n{content}"
        if char_count + len(entry_text) > max_chars and entries:
            break
        entries.append(entry_text)
        char_count += len(entry_text)
    entries.reverse()
    if len(entries) < len(debate_history):
        print(f"  [WARNING] debate history truncated: showing last {len(entries)} of {len(debate_history)} entries")
    truncated_note = f" (showing last {len(entries)} of {len(debate_history)} entries)" if len(entries) < len(debate_history) else ""
    context = f"\n\n--- DEBATE HISTORY{truncated_note} (retain and build on this) ---"
    context += "".join(entries)
    context += "\n--- END DEBATE HISTORY ---"
    return context

# ============================================================================
# NODE FUNCTIONS
# ============================================================================

# %% Intake agent system prompt (shared across intake, clarify, plan_topics)
INTAKE_SYSTEM_PROMPT = textwrap.dedent("""\
    You are a senior business strategy consultant conducting an intake session.
    Your job is to deeply understand a client's business question, frame it
    precisely and plan research topics.

    You retain full memory of the conversation so far. Build on prior exchanges
    rather than starting over.
""")

INTAKE_PARSE_PROMPT = textwrap.dedent("""\
    Parse the user's business question into structured fields: the core question,
    country/market, product idea, target customer, budget, time horizon, risk
    tolerance, and constraints. Extract whatever is mentioned; use 'Not specified'
    for missing fields.
""")

def _rebuild_chat_history(history: list[dict]) -> list:
    """Convert stored chat history dicts back to LangChain message objects."""
    msgs = []
    for h in history:
        role = h.get("role", "")
        content = h.get("content", "")
        if role == "system":
            msgs.append(SystemMessage(content=content))
        elif role == "human":
            msgs.append(HumanMessage(content=content))
        elif role == "ai":
            msgs.append(AIMessage(content=content))
        else:
            print(f"  [WARNING] Skipping history entry with unknown role: {role}")
    return msgs

def intake(state: dict) -> dict:
    print("\n")
    print("=" * 80)
    print("NODE: Intake — Parse Input")
    print("=" * 80)

    # Get query from config (terminal) or from messages (Studio)
    user_text = state.get("user_query", "")
    if not user_text:
        for m in reversed(state.get("messages", [])):
            if isinstance(m, HumanMessage):
                user_text = _msg_text(m)
                break
            if isinstance(m, dict) and m.get("type") == "human":
                user_text = m.get("content", "")
                break
            if hasattr(m, "type") and m.type == "human":
                user_text = _msg_text(m)
                break

    if not user_text:
        print("  -> No input found, skipping")
        return {}

    print(f"  Input: {user_text[:100]}...")

    # Parse structured fields
    try:
        structured = llm_intake.with_structured_output(IntakeOutput)
        result = structured.invoke([
            SystemMessage(content=INTAKE_PARSE_PROMPT),
            HumanMessage(content=user_text)])
        parsed = {
            "user_query": result.user_query,
            "country_or_market": result.country_or_market,
            "product_idea": result.product_idea,
            "target_customer": result.target_customer,
            "budget_range": result.budget_range,
            "time_horizon": result.time_horizon,
            "risk_tolerance": result.risk_tolerance,
            "constraints": result.constraints,
        }
    except Exception as e:
        print(f"  [Structured output failed: {e}] Using raw text as query")
        parsed = {"user_query": user_text}

    # Seed the intake conversation memory
    chat_history = [
        {"role": "system", "content": INTAKE_SYSTEM_PROMPT},
        {"role": "human", "content": user_text},
    ]

    print(f"  Parsed query: {parsed['user_query'][:100]}")
    return {**parsed, "intake_chat_history": chat_history}


# %% Node: Clarify Problem
CLARIFY_INSTRUCTION = textwrap.dedent("""\
    Based on our conversation so far, produce:
    1. A clear 2-3 sentence problem statement (approval-ready)
    2. Key constraints identified so far
    3. Questions for the user about unclear or missing aspects.
       You MUST always include at least 2-3 questions — ask about target
       customer, budget, timeline, product specifics, competitive positioning,
       or anything else that would sharpen the research. Never return an
       empty list.
""")

def clarify_problem(state: dict) -> dict:
    print("\n")
    print("=" * 80)
    print("NODE: Clarify Problem")
    print("=" * 80)

    # Build conversation from memory
    history = list(state.get("intake_chat_history", []))

    # If human gave feedback at Gate 1, add it to the conversation
    hf = state.get("human_feedback_1", "")
    if hf and hf != "approved":
        history.append({"role": "human", "content": hf})

    # Add the clarify instruction
    history.append({"role": "human", "content": CLARIFY_INSTRUCTION})

    # Call LLM with full conversation context
    chat_msgs = _rebuild_chat_history(history)
    try:
        structured = llm_intake.with_structured_output(ClarifyOutput)
        result = structured.invoke(chat_msgs)
        framing = result.problem_framing
        constraints_noted = result.constraints_noted
        questions = result.questions
    except Exception as e:
        print(f"  [Structured output failed: {e}] Falling back to text")
        resp = llm_intake.invoke(chat_msgs)
        framing = resp.content
        constraints_noted = ""
        questions = []

    print(f"  -> {len(questions)} questions generated")

    # Build display summary
    chat_summary = f"**Problem Framing:**\n{framing}"
    if constraints_noted:
        chat_summary += f"\n\n**Constraints:**\n{constraints_noted}"
    if questions:
        questions_text = "\n".join(f"  {i}. {q}" for i, q in enumerate(questions, 1))
        chat_summary += f"\n\n**Questions for you:**\n{questions_text}"

    # Update conversation memory — add the AI's response
    ai_response = f"Problem framing: {framing}"
    if constraints_noted:
        ai_response += f"\nConstraints: {constraints_noted}"
    if questions:
        ai_response += f"\nQuestions: {json.dumps(questions)}"
    history.append({"role": "ai", "content": ai_response})

    return {
        "problem_framing": framing,
        "constraints_noted": constraints_noted,
        "questions": questions,
        "intake_chat_history": history,
        "messages": [AIMessage(content=chat_summary)],
    }


# %% Node: Human Gate 1
def human_gate_1(state: dict) -> dict:
    """Pause for human review of problem framing."""
    print("\n")
    print("=" * 80)
    print("NODE: Human Gate 1 — Review Problem Framing")
    print("=" * 80)

    # Build the review content for the interrupt
    framing = state.get("problem_framing", "")
    questions = state.get("questions", [])

    constraints_noted = state.get("constraints_noted", "")

    review_content = f"--- PROBLEM FRAMING ---\n{framing}"
    if constraints_noted:
        review_content += f"\n\n--- CONSTRAINTS ---\n{constraints_noted}"
    if questions:
        questions_text = "\n".join(f"  {i}. {q}" for i, q in enumerate(questions, 1))
        review_content += f"\n\n--- QUESTIONS ---\n{questions_text}"

    review_content += (
        "\n\n---\n"
        "- Press Enter or 'approve' → proceed to research\n"
        "- Or answer the questions / provide feedback to revise"
    )

    if AUTO_APPROVE:
        print("  -> Auto-approved")
        return {"human_feedback_1": "approved"}

    feedback = interrupt(review_content)
    feedback = str(feedback).strip()

    if not feedback or feedback.lower() in ("approve", "approved", "ok", "yes", "skip", "looks good", "lgtm"):
        print("  -> Approved")
        return {"human_feedback_1": "approved"}

    clarify_round = state.get("clarify_round", 0) + 1
    if clarify_round >= MAX_CLARIFY_ROUNDS:
        print(f"  ** Max clarify rounds ({MAX_CLARIFY_ROUNDS}) reached. Forcing approve. **")
        return {"human_feedback_1": "approved", "clarify_round": clarify_round}

    print(f"  -> Feedback: {feedback[:100]} (round {clarify_round}/{MAX_CLARIFY_ROUNDS})")
    return {"human_feedback_1": feedback, "clarify_round": clarify_round}


# %% Routing: After Gate 1
def route_after_gate_1(state: dict) -> Literal["clarify_problem", "plan_research_topics"]:
    return "plan_research_topics" if state.get("human_feedback_1") == "approved" else "clarify_problem"


# %% Node: Plan Research Topics
PLAN_TOPICS_INSTRUCTION = textwrap.dedent("""\
    Based on our agreed problem framing, generate exactly {n}
    specific, searchable research topics. Each topic should be concrete enough
    for a researcher to investigate independently.
""")

def plan_research_topics(state: dict) -> dict:
    print("\n")
    print("=" * 80)
    print("NODE: Plan Research Topics")
    print("=" * 80)

    # Continue the intake conversation
    history = list(state.get("intake_chat_history", []))
    history.append({"role": "human", "content": PLAN_TOPICS_INSTRUCTION.format(n=MAX_RESEARCH_TOPICS)})

    chat_msgs = _rebuild_chat_history(history)
    structured = llm_intake.with_structured_output(TopicsOutput)
    result = structured.invoke(chat_msgs)
    topics = result.topics[:MAX_RESEARCH_TOPICS]

    print(f"  Topics ({len(topics)}):")
    for i, t in enumerate(topics, 1):
        print(f"    {i}. {t}")

    # Build the research brief — the clean handoff to the researcher agent
    topics_text = "\n".join(f"  {i}. {t}" for i, t in enumerate(topics, 1))
    research_brief = (
        f"## Problem\n{state.get('problem_framing', '') or state.get('user_query', '')}\n\n"
        f"## Context\n"
        f"- Market: {state.get('country_or_market', 'N/A')}\n"
        f"- Product: {state.get('product_idea', 'N/A')}\n"
        f"- Target customer: {state.get('target_customer', 'N/A')}\n"
        f"- Budget: {state.get('budget_range', 'N/A')}\n"
        f"- Time horizon: {state.get('time_horizon', 'N/A')}\n"
        f"- Risk tolerance: {state.get('risk_tolerance', 'N/A')}\n"
        f"- Constraints: {state.get('constraints', 'None')}\n"
        f"- Constraints noted: {state.get('constraints_noted', 'None')}\n\n"
        f"## Research Topics\n{topics_text}"
    )

    print(f"  Research brief: {len(research_brief)} chars")

    return {
        "research_topics": topics,
        "research_brief": research_brief,
        # Reset per-topic debate state
        "current_topic_idx": 0,
        "current_debate_round": 0,
        "current_topic_proposal": "",
        "current_topic_critique": "",
        "current_topic_limitations": [],
        "debate_converged": False,
        "debate_history": [],
        # Clear intake conversation memory (no longer needed)
        "intake_chat_history": [],
        "messages": [AIMessage(content=f"**Research Plan** ({len(topics)} topics):\n{topics_text}\n\nStarting per-topic debate loop...")],
    }


# %% Node: Research & Propose (NEW — replaces run_topic_research + proposer)
RESEARCH_PROPOSE_PROMPT = textwrap.dedent("""\
    You are a business research analyst and strategist. For the given topic:
    1. Research it thoroughly using any web results provided
    2. Produce specific findings with confidence levels
    3. Draft a strategic proposal for this topic
    4. Provide a one-line key recommendation

    You have access to the full debate history — all prior proposals, critic
    feedback, and human feedback for this topic. Build on and improve your
    prior work rather than starting from scratch.

    IMPORTANT: If a critic has reviewed your prior proposal, you MUST directly
    address every gap and concern they raised:
    - For each gap the critic identified, write an explicit response in the
      "gap_responses" field. One entry per gap, e.g.:
      "Gap: Lacks pricing data → Addressed: Added finding on SEK 25-45 price
       range from Euromonitor report" or
      "Gap: Missing competitor analysis → Cannot resolve: Private competitor
       financials are unavailable via public sources; recommend hiring a
       local market research firm"
    - Add NEW findings and data points, not just restate prior ones
    - If specific data cannot be found via web search, explicitly state this
      as a data gap and recommend how the client could obtain it
    - Do NOT repeat the same findings verbatim across rounds — each revision
      must contain materially new information or explicitly acknowledge limits
    - CRITICAL: If a gap has appeared in 2+ consecutive rounds and your web
      searches keep returning the same results, you MUST mark it as
      "Cannot resolve" in gap_responses. Do NOT claim "Addressed" if you
      are providing the same data you already provided in a prior round.
      Honest acknowledgment of limits is better than pretending old data is new.
""")

def research_and_propose(state: dict) -> dict:
    idx = state["current_topic_idx"]
    topics = state["research_topics"]
    topic = topics[idx]
    debate_round = state.get("current_debate_round", 0)

    print("\n")
    print("=" * 80)
    print(f"NODE: Research & Propose — Topic {idx+1}/{len(topics)}: {topic}")
    print(f"  Debate round: {debate_round}")
    print("=" * 80)

    # Web search — allowed on any round until budget exhausted
    # Web search — fresh budget each debate round
    # On revision rounds, search for the critic's specific gaps instead of the same generic query
    search_context = ""
    market = state.get("country_or_market", "")
    critique = state.get("current_topic_critique", "")

    if debate_round == 0 or not critique:
        # First round — broad topic search
        query = f"{topic} {market}" if market else topic
        try:
            results = web_search(query, max_results=MAX_WEB_SEARCH_CT)
        except Exception as e:
            print(f"  [WARNING] Web search failed: {e}")
            results = []
        if results:
            search_context = "\n\nWeb search results:\n" + "\n".join(
                f"- [{r['title']}]({r['url']}): {_truncate(r['content'], label='web search snippet')}" for r in results)
            print(f"  Found {len(results)} web sources (broad search)")
    else:
        # Revision rounds — targeted searches based on critic gaps
        # Extract gap keywords from critique to form targeted queries
        gap_queries = []
        for line in critique.split("\n"):
            line = line.strip()
            if line and len(line) > 20:
                # Use the gap text + market as a search query
                gap_query = f"{line[:100]} {market} {topic[:50]}" if market else f"{line[:100]} {topic[:50]}"
                gap_queries.append(gap_query)

        # Search for up to MAX_WEB_SEARCH_CT results across gap queries
        all_results = []
        per_gap = max(1, MAX_WEB_SEARCH_CT // max(len(gap_queries), 1))
        for gq in gap_queries[:MAX_WEB_SEARCH_CT]:
            try:
                results = web_search(gq, max_results=per_gap)
                all_results.extend(results)
            except Exception:
                pass
            if len(all_results) >= MAX_WEB_SEARCH_CT:
                break
        all_results = all_results[:MAX_WEB_SEARCH_CT]

        if all_results:
            search_context = "\n\nWeb search results (targeted at critic gaps):\n" + "\n".join(
                f"- [{r['title']}]({r['url']}): {_truncate(r['content'], label='web search snippet')}" for r in all_results)
            print(f"  Found {len(all_results)} web sources (targeted gap search)")

    # Build debate history context (researcher ⇄ critic + human, truncated from front)
    debate_history = state.get("debate_history", [])
    debate_context = _build_debate_context(debate_history)

    research_brief = state.get("research_brief", "")
    user_msg = f"""Research this topic: {topic}

{research_brief}
{search_context}{debate_context}"""

    try:
        structured = llm_researcher.with_structured_output(TopicProposalOutput)
        result = structured.invoke([SystemMessage(content=RESEARCH_PROPOSE_PROMPT),
                                    HumanMessage(content=user_msg)])
        proposal_dict = result.model_dump()
        # Serialize findings for storage
        proposal_dict["findings"] = [f.model_dump() if hasattr(f, "model_dump") else f for f in result.findings]
        gap_responses_text = ""
        if result.gap_responses:
            gap_responses_text = (
                f"\n\n**Gap Responses ({len(result.gap_responses)}):**\n"
                + "\n".join(f"- {r}" for r in result.gap_responses)
            )
        proposal_text = (
            f"**Topic: {result.topic}**\n\n"
            f"**Summary:** {result.summary}\n\n"
            f"**Proposal:** {result.proposal}\n\n"
            f"**Key Recommendation:** {result.key_recommendation}\n\n"
            f"**Findings ({len(result.findings)}):**\n"
            + "\n".join(f"- [{f.confidence}] {f.claim}" for f in result.findings)
            + gap_responses_text
        )
    except Exception as e:
        print(f"  [Structured output failed: {e}] Falling back to text")
        resp = llm_researcher.invoke([SystemMessage(content=RESEARCH_PROPOSE_PROMPT),
                           HumanMessage(content=user_msg)])
        proposal_text = resp.content
        proposal_dict = {"topic": topic, "findings": [], "summary": _truncate(resp.content, label="proposal fallback summary"),
                         "proposal": resp.content, "key_recommendation": "See proposal text"}

    print(f"  -> Proposal generated ({len(proposal_text)} chars)")
    print()
    print(f"  Summary: {proposal_dict.get('summary', 'N/A')}")
    print(f"  Key Recommendation: {proposal_dict.get('key_recommendation', 'N/A')}")
    findings_list = proposal_dict.get("findings", [])
    if findings_list:
        print(f"  Findings ({len(findings_list)}):")
        for f in findings_list[:5]:
            claim = f.get("claim", "") if isinstance(f, dict) else str(f)
            conf = f.get("confidence", "?") if isinstance(f, dict) else "?"
            print(f"    [{conf}] {claim}")
    gap_responses_list = proposal_dict.get("gap_responses", [])
    if gap_responses_list:
        print(f"  Gap Responses ({len(gap_responses_list)}):")
        for r in gap_responses_list:
            print(f"    -> {r}")

    # Append to debate history
    new_debate_history = list(debate_history) + [
        {"role": "researcher", "content": proposal_text}
    ]

    return {
        "current_topic_proposal": json.dumps(proposal_dict, default=str),
        "debate_converged": False,
        "debate_history": new_debate_history,
        "messages": [AIMessage(content=f"**[Topic {idx+1}/{len(topics)} — Round {debate_round+1}]**\n\n{proposal_text}")],
    }


# %% Node: Topic Critic (NEW)
TOPIC_CRITIC_PROMPT = textwrap.dedent("""\
    You are a demanding business strategy critic reviewing a per-topic proposal.
    Evaluate the proposal's quality, identify gaps, and decide if it's ready
    for human review.

    You have access to the full debate history — all prior proposals, your own
    prior critiques, and any human feedback. Use this context to track whether
    the researcher has addressed previous concerns and improved over rounds.

    CRITICAL: On the FIRST round (round 0), you must ALWAYS set converged=false.
    Your job is to push the researcher to improve. Identify specific, actionable
    gaps — missing data points, weak evidence, unsupported claims, missing
    competitor analysis, vague recommendations, etc. Be demanding.

    On subsequent rounds (round 1+), set converged=true if the researcher has
    meaningfully addressed your prior gaps and the proposal is now substantive
    and well-supported. Set converged=false if there are still significant gaps
    the researcher can realistically fill with additional web research.

    IMPORTANT convergence rules:
    - If the researcher has explicitly stated that certain data is unavailable
      via web search and recommended primary research methods, accept that —
      do NOT keep requesting the same unfillable gap across rounds
    - CRITICAL: If you are requesting the SAME gap (or substantially similar
      wording) for a 2nd time, you MUST set converged=true. The researcher
      has already tried and cannot find this data. Continuing to ask for it
      wastes rounds. Move the unresolved gap to limitations and converge.
    - Look at the debate history — if you see the same gap appearing in your
      prior critiques, that is your signal to converge, not to ask again.
    - In the first half of the debate (before round {half_max}), you MAY raise
      new gaps you missed before — being thorough is more important than being consistent
    - In the second half (round {half_max}+), do NOT raise new gaps that weren't
      in your previous critique unless the revision introduced new problems
    - Your job is to improve the proposal, not to demand perfection
    - When converged=true: gaps MUST be empty, revision_guidance MUST be
      empty. Any remaining concerns that can't be resolved via web research
      go in the "limitations" field instead.
    - Only populate the "limitations" field with items the RESEARCHER has
      explicitly stated cannot be resolved via web search or further analysis
      (look for their gap_responses). Do NOT invent limitations on your own.
      If the researcher hasn't flagged anything as unresolvable, return an
      empty limitations list.
""")

def topic_critic(state: dict) -> dict:
    idx = state["current_topic_idx"]
    topics = state["research_topics"]
    topic = topics[idx]
    debate_round = state.get("current_debate_round", 0)

    print("\n")
    print("=" * 80)
    print(f"NODE: Topic Critic — Topic {idx+1}/{len(topics)}: {topic}")
    print(f"  Debate round: {debate_round}")
    print("=" * 80)

    proposal = state.get("current_topic_proposal", "")

    # Build debate history context (researcher ⇄ critic + human, truncated from front)
    debate_history = state.get("debate_history", [])
    debate_context = _build_debate_context(debate_history)

    research_brief = state.get("research_brief", "")
    user_msg = f"""Topic: {topic}

{research_brief}

Current proposal:
{_truncate(proposal, label="proposal for critic")}

Debate round: {debate_round + 1} of max {MAX_DEBATE_ROUNDS}
{debate_context}"""

    try:
        structured = llm_critic.with_structured_output(CriticOutput)
        critic_prompt = TOPIC_CRITIC_PROMPT.format(half_max=MAX_DEBATE_ROUNDS // 2)
        result = structured.invoke([SystemMessage(content=critic_prompt),
                                    HumanMessage(content=user_msg)])
        assessment = result.assessment
        gaps = result.gaps
        converged = result.converged
        revision_guidance = result.revision_guidance
        limitations = result.limitations
    except Exception as e:
        print(f"  [Structured output failed: {e}] Assuming converged")
        assessment = "Unable to parse critic output; treating as converged."
        gaps = []
        converged = True
        revision_guidance = ""
        limitations = []

    # Build display text — show gaps/guidance only if not converged, limitations only if converged
    critique_text = f"**Assessment:** {assessment}\n\n**Converged:** {'Yes' if converged else 'No'}"
    if not converged:
        critique_text += f"\n\n**Gaps:** {', '.join(gaps) if gaps else 'None'}"
        if revision_guidance:
            critique_text += f"\n\n**Revision guidance:** {revision_guidance}"
    if converged and limitations:
        critique_text += f"\n\n**Limitations:** {', '.join(limitations)}"

    # Terminal output
    print(f"  -> Converged: {converged}")
    print(f"  Assessment: {assessment}")
    if not converged:
        print(f"  Gaps: {len(gaps) if gaps else 'None'}")
        for g in gaps:
            print(f"    - {g}")
        print(f"  Guidance: {revision_guidance if revision_guidance else 'None'}")
    if converged and limitations:
        print(f"  Limitations: {len(limitations)}")
        for l in limitations:
            print(f"    - {l}")

    # Append to debate history
    new_debate_history = list(debate_history) + [
        {"role": "critic", "content": critique_text}
    ]

    # Store critique — include limitations when converged
    if converged:
        stored_critique = f"{assessment}\n\nLimitations: {', '.join(limitations) if limitations else 'None'}"
    else:
        stored_critique = f"{assessment}\n\nGaps: {', '.join(gaps) if gaps else 'None'}\n\nGuidance: {revision_guidance if revision_guidance else 'None'}"

    return {
        "current_topic_critique": stored_critique,
        "current_topic_limitations": limitations,
        "debate_converged": converged,
        "current_debate_round": debate_round + 1,
        "debate_history": new_debate_history,
        "messages": [AIMessage(content=f"**[Critic — Topic {idx+1}/{len(topics)}]**\n\n{critique_text}")],
    }


# %% Routing: After Critic
def route_after_critic(state: dict) -> Literal["human_gate_2", "research_and_propose"]:
    """If converged or max debate rounds reached → human gate. Else → revise."""
    if state.get("debate_converged", False):
        return "human_gate_2"
    if state.get("current_debate_round", 0) >= MAX_DEBATE_ROUNDS:
        print(f"  [Max debate rounds ({MAX_DEBATE_ROUNDS}) reached — moving to human gate]")
        return "human_gate_2"
    return "research_and_propose"


def _build_approved_entry(state: dict, topic: str) -> dict:
    """Build an approved topic entry from current state. Used by both approve paths."""
    try:
        proposal_data = json.loads(state.get("current_topic_proposal", "{}"))
    except (json.JSONDecodeError, TypeError):
        proposal_data = {"topic": topic, "proposal": state.get("current_topic_proposal", "")}
    return {
        "topic": topic,
        "proposal": proposal_data.get("proposal", ""),
        "key_recommendation": proposal_data.get("key_recommendation", ""),
        "summary": proposal_data.get("summary", ""),
        "findings": proposal_data.get("findings", []),
        "critic_assessment": state.get("current_topic_critique", ""),
        "limitations": state.get("current_topic_limitations", []),
        "debate_rounds": state.get("current_debate_round", 1),
        "debate_converged": state.get("debate_converged", False),
    }

_GATE2_RESET = {
    "current_debate_round": 0,
    "current_topic_proposal": "",
    "current_topic_critique": "",
    "current_topic_limitations": [],
    "proposal_revision_round": 0,
    "debate_converged": False,
    "debate_history": [],
}


# %% Node: Human Gate 2 (per-topic)
def human_gate_2(state: dict) -> dict:
    """Per-topic human review. Approve → append to approved_topics + advance.
    Revise → loop back to research_and_propose for same topic."""
    idx = state["current_topic_idx"]
    topics = state["research_topics"]
    topic = topics[idx]

    print("\n")
    print("=" * 80)
    print(f"NODE: Human Gate 2 — Review Topic {idx+1}/{len(topics)}: {topic}")
    print("=" * 80)

    # Parse proposal for display
    print()
    proposal_raw = state.get("current_topic_proposal", "")
    try:
        proposal_data = json.loads(proposal_raw)
        proposal_display = (
            f"Summary: {proposal_data.get('summary', 'N/A')}\n\n"
            f"Proposal: {proposal_data.get('proposal', 'N/A')}\n\n"
            f"Key Recommendation: {proposal_data.get('key_recommendation', 'N/A')}"
        )
    except (json.JSONDecodeError, TypeError):
        proposal_display = _truncate(proposal_raw, label="proposal display")

    critique_display = state.get("current_topic_critique", "No critique available")
    debate_rounds = state.get("current_debate_round", 0)
    converged = state.get("debate_converged", False)
    converge_status = "Converged" if converged else f"Not converged (max {MAX_DEBATE_ROUNDS} rounds reached)"

    if AUTO_APPROVE:
        print("  -> Auto-approved")
        feedback = "approved"
    else:
        limitations = state.get("current_topic_limitations", [])
        limitations_note = ""
        if limitations:
            limitations_note = (
                "\n\n--- LIMITATIONS (cannot be resolved via web research) ---\n"
                + "\n".join(f"  - {l}" for l in limitations)
            )

        feedback = interrupt(
            f"Topic {idx+1}/{len(topics)}: '{topic}' ({debate_rounds} debate rounds, {converge_status})\n\n"
            f"--- PROPOSAL ---\n{proposal_display}\n\n"
            f"--- CRITIC ---\n{critique_display}"
            f"{limitations_note}\n\n"
            "---\n"
            "- Press Enter or 'approve' → accept and move to next topic\n"
            "- Or provide feedback to revise this topic's proposal\n"
            "  (Note: limitations listed above require primary research and cannot be addressed by further web search)"
        )
        feedback = str(feedback).strip()

    is_approved = not feedback or feedback.lower() in (
        "approve", "approved", "ok", "yes", "skip", "looks good", "lgtm"
    )

    if is_approved:
        print(f"  -> Approved topic {idx+1}")
        return {
            "human_feedback_2": "approved",
            "approved_topics": [_build_approved_entry(state, topic)],
            "current_topic_idx": idx + 1,
            **_GATE2_RESET,
        }
    else:
        revision_round = state.get("proposal_revision_round", 0) + 1
        if revision_round >= MAX_HUMAN_REVISION_ON_PROPOSAL:
            print(f"  ** Max human revisions ({MAX_HUMAN_REVISION_ON_PROPOSAL}) reached. Forcing approve. **")
            return {
                "human_feedback_2": "approved",
                "approved_topics": [_build_approved_entry(state, topic)],
                "current_topic_idx": idx + 1,
                **_GATE2_RESET,
            }

        print(f"  -> Revise requested: {feedback[:100]} (revision {revision_round}/{MAX_HUMAN_REVISION_ON_PROPOSAL})")

        # Structure human feedback as critic-style gaps so the researcher
        # treats it with the same rigor as critic gaps
        human_critique = (
            f"**Assessment:** Human reviewer requested revisions.\n\n"
            f"**Gaps:**\n- {feedback}\n\n"
            f"**Revision guidance:** Address the human reviewer's feedback above."
        )

        # Append human feedback to debate history so both agents see it
        debate_history = list(state.get("debate_history", []))
        debate_history.append({"role": "human", "content": human_critique})

        return {
            "human_feedback_2": feedback,
            "current_topic_critique": human_critique,
            "proposal_revision_round": revision_round,
            "current_debate_round": 0,  # Reset so human feedback gets a full debate cycle
            "debate_converged": False,
            "debate_history": debate_history,
        }


# %% Routing: After Gate 2
def route_after_gate_2(state: dict) -> Literal["research_and_propose", "synthesizer"]:
    """After gate 2:
    - If revise requested → research_and_propose (same topic, idx not advanced)
    - If approved + more topics → research_and_propose (next topic, idx advanced)
    - If approved + all done → synthesizer
    """
    if state.get("human_feedback_2", "") not in ("approved", ""):
        # Revise — same topic
        return "research_and_propose"

    # Approved — check if more topics
    idx = state.get("current_topic_idx", 0)
    total = len(state.get("research_topics", []))
    if idx < total:
        return "research_and_propose"
    return "synthesizer"


# %% Node: Synthesizer
SYNTH_PROMPT = textwrap.dedent("""\
    You are a senior business consultant writing a final recommendation report.
    You have approved proposals for multiple research topics. Synthesize them into:
    1. **Executive Summary** (5 bullet points)
    2. **Problem Framing & Assumptions**
    3. **Evidence Snapshot** (by topic, with key findings and citations)
    4. **Strategic Recommendations** (one per topic, with tradeoffs)
    5. **Integrated Strategy & Rationale** (how the pieces fit together)
    6. **Risks & Mitigations**
    7. **Known Limitations & Data Gaps** (from critic reviews — what couldn't be verified)
    8. **References** — list ALL source URLs from the research findings. Format each
       as a numbered entry: [N] Title — URL. These are the actual web sources the
       researcher found. In sections 3 and 4, cite sources using [N] notation that
       corresponds to this reference list.
    Be specific, evidence-based, actionable. Use markdown.
    Each topic entry includes a "limitations" field listing data gaps the critic
    identified as unresolvable via web research. Surface these honestly in section 7
    and factor them into your risk assessment in section 6.
    Topics also have "debate_converged" — if false, the debate hit the max round
    limit without full agreement, so treat those findings with extra caution.
    Each finding has a "sources" list with "title" and "url" — use these for the
    References section. Do NOT invent or hallucinate URLs.
""")

def synthesizer(state: dict) -> dict:
    print("\n")
    print("=" * 80)
    print("NODE: Synthesize Recommendation (runs once)")
    print("=" * 80)

    approved = state.get("approved_topics", [])
    print(f"  Synthesizing {len(approved)} approved topics")

    approved_text = _truncate(json.dumps(approved, indent=2, default=str), label="synthesizer approved_topics")
    research_brief = state.get("research_brief", "")

    user_msg = _truncate(f"""{research_brief}

Approved topic proposals:
{approved_text}

Write a comprehensive recommendation report with all 8 sections.
IMPORTANT: Include a References section at the end with all source URLs from the findings. Cite them using [N] notation throughout the report.""", label="synthesizer input")

    resp = llm_synthesizer.invoke([SystemMessage(content=SYNTH_PROMPT),
                                    HumanMessage(content=user_msg)])
    recommendation = resp.content

    if not recommendation:
        print("  [WARNING] Empty response, retrying with shorter prompt...")
        resp = llm_synthesizer.invoke([SystemMessage(content=SYNTH_PROMPT),
                                        HumanMessage(content=f"Problem: {state.get('problem_framing', '')}\nTopics: {_truncate(approved_text, label='synthesizer retry')}\nWrite recommendation.")])
        recommendation = resp.content

    print(f"  -> Report generated ({len(recommendation)} chars)")

    return {
        "recommendation": recommendation,
        "messages": [AIMessage(content=f"**Recommendation Report:**\n\n{recommendation}")],
    }


# %% Node: Action Plan 90-Day
ACTION_PLAN_PROMPT = textwrap.dedent("""\
    You are a business execution planner. Create a 90-day action plan:
    ### Days 1-30: Foundation & Validation (3-5 actions, milestones, KPIs)
    ### Days 31-60: Build & Test (3-5 actions, milestones, KPIs)
    ### Days 61-90: Launch & Measure (3-5 actions, milestones, KPIs)
    ### Data Gaps & Next Steps
    The recommendation report may include known limitations and data gaps from
    the research phase. Incorporate specific actions to address these gaps
    (e.g., commission a market study, run a focus group, hire a local consultant).
    Be specific to the business context. Use markdown.
""")

def action_plan_90d(state: dict) -> dict:
    print("\n")
    print("=" * 80)
    print("NODE: 90-Day Action Plan")
    print("=" * 80)

    n_topics = len(state.get("approved_topics", []))
    topic_names = [t.get("topic", "?") for t in state.get("approved_topics", [])]

    # Build conversation history for action plan
    action_plan_history = list(state.get("action_plan_history", []))

    if not action_plan_history:
        # First run — seed with system prompt + synthesizer output + context
        action_plan_history.append({"role": "system", "content": ACTION_PLAN_PROMPT})
        action_plan_history.append({"role": "human", "content":
            f"""Recommendation report:
{state.get('recommendation', 'N/A')}

Market: {state.get('country_or_market', 'N/A')}
Product: {state.get('product_idea', 'N/A')}
Budget: {state.get('budget_range', 'N/A')}
Timeline: {state.get('time_horizon', 'N/A')}
Research topics covered: {', '.join(topic_names)}

Create a 90-day action plan based on the recommendation report above."""})
    else:
        # Gate 3 sent us back — human feedback already appended by human_gate_3
        action_plan_history.append({"role": "human", "content":
            "Revise the action plan based on the feedback above."})

    chat_msgs = _rebuild_chat_history(action_plan_history)
    resp = llm_synthesizer.invoke(chat_msgs)
    plan = resp.content

    # Append plan output to history
    action_plan_history.append({"role": "ai", "content": plan})

    final = f"""{'='*60}
FINAL MBA STRATEGY REPORT (v4)
{'='*60}

{state.get('recommendation', '[No recommendation]')}

{'='*60}
90-DAY ACTION PLAN
{'='*60}

{plan}

{'='*60}
METADATA
{'='*60}
Agents: {json.dumps({k: v['model'] for k, v in AGENTS.items()})}
Topics researched: {n_topics}
Topic names: {', '.join(topic_names)}
Web search: Tavily
"""

    print(f"  -> Action plan generated ({len(plan)} chars)")

    return {
        "action_plan": plan,
        "action_plan_history": action_plan_history,
        "final_output": final,
        "status": "pending_final_approval",
        "messages": [AIMessage(content=f"**90-Day Action Plan:**\n\n{plan}")],
    }


# %% Node: Human Gate 3 (NEW — final approval)
def human_gate_3(state: dict) -> dict:
    """Final human approval of the synthesized report + action plan."""
    print("\n")
    print("=" * 80)
    print("NODE: Human Gate 3 — Final Approval")
    print("=" * 80)

    recommendation = state.get("recommendation", "")
    action_plan = state.get("action_plan", "")

    if AUTO_APPROVE:
        print("  -> Auto-approved")
        return {"human_feedback_3": "approved", "status": "finalize"}

    feedback = interrupt(
        f"--- RECOMMENDATION REPORT ---\n{recommendation}\n\n"
        f"--- 90-DAY ACTION PLAN ---\n{action_plan}\n\n"
        "---\n"
        "- Press Enter or 'approve' → finalize and save the report\n"
        "- Or provide feedback to revise the action plan"
    )
    feedback = str(feedback).strip()

    if not feedback or feedback.lower() in ("approve", "approved", "ok", "yes", "skip", "looks good", "lgtm"):
        print("  -> Approved — finalizing")
        return {"human_feedback_3": "approved", "status": "finalize"}

    revision_round = state.get("plan_revision_round", 0) + 1
    if revision_round >= MAX_HUMAN_REVISION_ON_PLAN:
        print(f"  ** Max plan revisions ({MAX_HUMAN_REVISION_ON_PLAN}) reached. Forcing approve. **")
        return {"human_feedback_3": "approved", "status": "finalize", "plan_revision_round": revision_round}

    print(f"  -> Feedback: {feedback[:100]} (revision {revision_round}/{MAX_HUMAN_REVISION_ON_PLAN})")

    # Append human feedback to action plan history so it retains the conversation
    action_plan_history = list(state.get("action_plan_history", []))
    action_plan_history.append({"role": "human", "content": feedback})

    return {"human_feedback_3": feedback, "plan_revision_round": revision_round, "action_plan_history": action_plan_history}


# %% Routing: After Gate 3
def route_after_gate_3(state: dict) -> Literal["action_plan_90d", "__end__"]:
    """Approved → END. Feedback → revise action plan."""
    if state.get("human_feedback_3") == "approved":
        return "__end__"
    return "action_plan_90d"


print("All node functions defined (11 nodes)")

# ============================================================================
# BUILD GRAPH
# ============================================================================

# %% Build Graph
def build_graph(checkpointer=None):
    g = StateGraph(MBAAgentState)

    # Add all 11 nodes
    g.add_node("intake", intake)
    g.add_node("clarify_problem", clarify_problem)
    g.add_node("human_gate_1", human_gate_1)
    g.add_node("plan_research_topics", plan_research_topics)
    g.add_node("research_and_propose", research_and_propose)
    g.add_node("topic_critic", topic_critic)
    g.add_node("human_gate_2", human_gate_2)
    g.add_node("synthesizer", synthesizer)
    g.add_node("action_plan_90d", action_plan_90d)
    g.add_node("human_gate_3", human_gate_3)

    # ── Phase 1: Intake ──
    g.add_edge(START, "intake")
    g.add_edge("intake", "clarify_problem")
    g.add_edge("clarify_problem", "human_gate_1")
    g.add_conditional_edges("human_gate_1", route_after_gate_1, {
        "clarify_problem": "clarify_problem",
        "plan_research_topics": "plan_research_topics",
    })

    # ── Phase 2: Per-Topic Debate Loop ──
    g.add_edge("plan_research_topics", "research_and_propose")
    g.add_edge("research_and_propose", "topic_critic")
    g.add_conditional_edges("topic_critic", route_after_critic, {
        "human_gate_2": "human_gate_2",
        "research_and_propose": "research_and_propose",
    })
    g.add_conditional_edges("human_gate_2", route_after_gate_2, {
        "research_and_propose": "research_and_propose",
        "synthesizer": "synthesizer",
    })

    # ── Phase 3: Synthesis & Final Approval ──
    g.add_edge("synthesizer", "action_plan_90d")
    g.add_edge("action_plan_90d", "human_gate_3")
    g.add_conditional_edges("human_gate_3", route_after_gate_3, {
        "action_plan_90d": "action_plan_90d",
        "__end__": END,
    })

    return g.compile(checkpointer=checkpointer)


# Module-level graph for `langgraph dev` / Studio (server provides checkpointer)
graph = build_graph()

print("=" * 80)
print("GRAPH BUILT — v4")
print("=" * 80)
print(f"  Nodes: {len(graph.get_graph().nodes)}")
print(f"  Edges: {len(graph.get_graph().edges)}")

# ============================================================================
# TERMINAL RUNNER
# ============================================================================

# %% Run Agent
if __name__ == "__main__":
    from langgraph.checkpoint.memory import MemorySaver
    from langgraph.types import Command

    print("=" * 80)
    print("RUNNING MBA STRATEGY AGENT v4 (terminal mode)")
    print("=" * 80)
    print(f"  Query:        {INPUT_QUERY}")
    print(f"  Auto-approve: {AUTO_APPROVE}")
    print()
    print("  Agents:")
    for name, cfg in AGENTS.items():
        temp_str = f"  temp={cfg['temperature']}" if "temperature" in cfg else ""
        print(f"    {name:12s}  {cfg['model']:40s}{temp_str}")
    print("=" * 80)

    agent = build_graph(checkpointer=MemorySaver())
    config = {"configurable": {"thread_id": "mba-v4-demo-1"}}

    # Only user_query is required — all other fields use .get() with defaults
    initial_state = {"user_query": INPUT_QUERY}

    # First invoke — runs until first interrupt (or completion if AUTO_APPROVE)
    result = agent.invoke(initial_state, config)

    # Interactive loop: handle interrupts until graph completes
    while True:
        snapshot = agent.get_state(config)
        if not snapshot.next:
            break  # graph completed — no pending nodes

        # Show interrupt prompt (contains full review content)
        for task in snapshot.tasks:
            if hasattr(task, 'interrupts'):
                for intr in task.interrupts:
                    print("\n" + "=" * 80)
                    print(intr.value)
                    print("=" * 80)

        # Get human input
        feedback = input("\n> ").strip()
        if not feedback:
            feedback = "approved"

        # Resume graph from interrupt
        result = agent.invoke(Command(resume=feedback), config)

    # ── Display Results ──
    print("\n" + "=" * 80)
    print("AGENT COMPLETE")
    print("=" * 80)
    print(result.get("final_output", "No output generated."))

    # ── Save Results ──
    print("=" * 80)
    print("SAVING RESULTS")
    print("=" * 80)

    output_dir = Path(__file__).resolve().parent / "results"
    output_dir.mkdir(exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")

    json_path = output_dir / f"mba_v4_output_{ts}.json"
    serializable = {}
    for k, v in result.items():
        if k == "messages":
            serializable[k] = [
                {"role": getattr(m, "type", "unknown"), "content": getattr(m, "content", str(m))}
                for m in v
            ]
        elif isinstance(v, (str, int, float, bool, list, dict, type(None))):
            serializable[k] = v

    with open(json_path, "w") as f:
        json.dump(serializable, f, indent=2, ensure_ascii=False)
    print(f"  JSON:     {json_path}")

    md_path = output_dir / f"mba_v4_report_{ts}.md"
    with open(md_path, "w") as f:
        f.write(f"# MBA Strategy Report v4\n\n")
        f.write(f"**Generated:** {datetime.now().isoformat()}\n")
        agents_str = ', '.join(f'{k}={v["model"]}' for k, v in AGENTS.items())
        f.write(f"**Agents:** {agents_str}\n\n---\n\n")
        f.write(result.get("recommendation", ""))
        f.write(f"\n\n---\n\n")
        f.write(result.get("action_plan", ""))
    print(f"  Markdown: {md_path}")

    # ── DOCX output ──
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import re

        doc = Document()

        # -- Styles --
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        # -- Title page --
        title = doc.add_heading("MBA Strategy Report", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y at %H:%M')}")
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        agents_para = doc.add_paragraph()
        agents_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = agents_para.add_run(f"Models: {agents_str}")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        doc.add_page_break()

        # -- Helper: convert markdown text to docx paragraphs --
        def _md_to_docx(doc, md_text):
            """Simple markdown-to-docx converter for headings, bold, bullets, and paragraphs."""
            lines = md_text.split("\n")
            i = 0
            while i < len(lines):
                line = lines[i]
                stripped = line.strip()

                # Skip empty lines
                if not stripped:
                    i += 1
                    continue

                # Headings
                if stripped.startswith("######"):
                    doc.add_heading(stripped.lstrip("#").strip(), level=5)
                elif stripped.startswith("#####"):
                    doc.add_heading(stripped.lstrip("#").strip(), level=4)
                elif stripped.startswith("####"):
                    doc.add_heading(stripped.lstrip("#").strip(), level=4)
                elif stripped.startswith("###"):
                    doc.add_heading(stripped.lstrip("#").strip(), level=3)
                elif stripped.startswith("##"):
                    doc.add_heading(stripped.lstrip("#").strip(), level=2)
                elif stripped.startswith("#"):
                    doc.add_heading(stripped.lstrip("#").strip(), level=1)
                # Bullet points
                elif stripped.startswith("- ") or stripped.startswith("* "):
                    text = stripped[2:]
                    p = doc.add_paragraph(style="List Bullet")
                    _add_formatted_text(p, text)
                # Numbered lists
                elif re.match(r"^\d+\.\s", stripped):
                    text = re.sub(r"^\d+\.\s", "", stripped)
                    p = doc.add_paragraph(style="List Number")
                    _add_formatted_text(p, text)
                # Regular paragraph
                else:
                    p = doc.add_paragraph()
                    _add_formatted_text(p, stripped)

                i += 1

        def _add_formatted_text(paragraph, text):
            """Add text with bold (**text**) formatting support."""
            parts = re.split(r"(\*\*[^*]+\*\*)", text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = paragraph.add_run(part[2:-2])
                    run.bold = True
                else:
                    paragraph.add_run(part)

        # -- Recommendation Report --
        recommendation = result.get("recommendation", "")
        if recommendation:
            _md_to_docx(doc, recommendation)

        doc.add_page_break()

        # -- Action Plan --
        action_plan = result.get("action_plan", "")
        if action_plan:
            _md_to_docx(doc, action_plan)

        docx_path = output_dir / f"mba_v4_report_{ts}.docx"
        doc.save(str(docx_path))
        print(f"  DOCX:     {docx_path}")

    except ImportError:
        print("  [SKIP] python-docx not installed — run: pip install python-docx")
    except Exception as e:
        print(f"  [WARNING] DOCX generation failed: {e}")

    print("=" * 80)
    print("COMPLETE")
    print("=" * 80)
